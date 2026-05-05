"""
Google Drive service for Alfred.
Fetches files from Drive API and caches them in SQLite drive_index table.
First call is slow; subsequent calls read from local index (fast).
Index refreshes automatically if older than 30 minutes.
Supports shared drives (supportsAllDrives / includeItemsFromAllDrives).
Per-user DB support: pass user_conn to use per-user DB; falls back to shared DB.
"""
import httpx
import sqlite3 as _sq
_SHARED_DB = '/opt/alfred/data/alfred.db'
def _shared_conn():
    return _sq.connect(_SHARED_DB)

from datetime import datetime, timedelta

DRIVE_API = "https://www.googleapis.com/drive/v3"

ICON_MAP = {
    "application/pdf": "PDF",
    "application/vnd.google-apps.document": "Google 文件",
    "application/vnd.google-apps.spreadsheet": "試算表",
    "application/vnd.google-apps.presentation": "簡報",
    "application/vnd.google-apps.folder": "資料夾",
    "image/jpeg": "圖片", "image/png": "圖片", "image/gif": "圖片",
    "video/mp4": "影片", "audio/mpeg": "音樂",
    "application/zip": "壓縮檔",
    "text/plain": "文字檔",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "Word",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "Excel",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation": "PowerPoint",
    "application/msword": "Word",
}

def _get_conn(user_conn=None):
    """Return (conn, should_close). If user_conn provided, use it (caller manages lifecycle)."""
    if user_conn is not None:
        return user_conn, False
    c = _shared_conn()
    return c, True

def _token(db_func) -> str | None:
    try:
        import gcal_service
        return gcal_service._get_access_token(db_func)
    except Exception:
        return None

def _type_label(mime: str) -> str:
    return ICON_MAP.get(mime, "檔案")

def _fmt_size(b) -> str:
    try:
        b = int(b)
        if b < 1024: return f"{b}B"
        if b < 1048576: return f"{b//1024}KB"
        return f"{b/1048576:.1f}MB"
    except Exception:
        return "—"

def _ensure_drive_index_table(conn):
    """Ensure drive_index table and drive_name column exist in the given conn."""
    conn.execute("""CREATE TABLE IF NOT EXISTS drive_index
        (id TEXT PRIMARY KEY, name TEXT, mime_type TEXT, size TEXT,
         modified TEXT, url TEXT, drive_name TEXT, indexed_at TEXT)""")
    cols = [row[1] for row in conn.execute("PRAGMA table_info(drive_index)").fetchall()]
    if "drive_name" not in cols:
        conn.execute("ALTER TABLE drive_index ADD COLUMN drive_name TEXT DEFAULT ''")
    conn.commit()

def _index_is_fresh(db_func, max_age_minutes: int = 30, user_conn=None) -> bool:
    """Check if index was updated within max_age_minutes."""
    try:
        c, should_close = _get_conn(user_conn)
        _ensure_drive_index_table(c)
        row = c.execute("SELECT indexed_at FROM drive_index ORDER BY indexed_at DESC LIMIT 1").fetchone()
        if should_close:
            c.close()
        if not row:
            return False
        last = datetime.fromisoformat(row[0])
        return datetime.now() - last < timedelta(minutes=max_age_minutes)
    except Exception:
        return False

def _ensure_drive_name_column():
    """Add drive_name column to drive_index if it does not exist yet (shared DB, kept for backward compat)."""
    try:
        c = _shared_conn()
        _ensure_drive_index_table(c)
        c.close()
    except Exception:
        pass

def _save_to_index(db_func, files: list[dict], user_conn=None):
    """Upsert Drive file metadata into local SQLite index."""
    c, should_close = _get_conn(user_conn)
    _ensure_drive_index_table(c)
    now = datetime.now().isoformat()
    for f in files:
        c.execute(
            """INSERT OR REPLACE INTO drive_index (id,name,mime_type,size,modified,url,indexed_at,drive_name)
               VALUES (?,?,?,?,?,?,?,?)""",
            (f["id"], f["name"], f.get("mimeType",""), _fmt_size(f.get("size",0)),
             f.get("modifiedTime","")[:10], f.get("webViewLink",""), now,
             f.get("drive_name","我的雲端硬碟"))
        )
    c.commit()
    if should_close:
        c.close()

def _fetch_shared_drives(token: str) -> list[dict]:
    """列出所有共用雲端硬碟，回傳 [{id, name}]。"""
    try:
        r = httpx.get(f"{DRIVE_API}/drives",
            headers={"Authorization": f"Bearer {token}"},
            params={"pageSize": 100, "fields": "drives(id,name)"},
            timeout=15)
        drives = r.json().get("drives", [])
        print(f"[drive_service] 共用雲端硬碟數量：{len(drives)}")
        return drives
    except Exception as e:
        print(f"[drive_service] _fetch_shared_drives 失敗：{e}")
        return []

def _fetch_from_api(db_func, query: str = "", max_results: int = 100, user_conn=None) -> list[dict]:
    """Fetch from Drive API (含共用雲端硬碟) and save to index. Returns index-format dicts."""
    token = _token(db_func)
    if not token:
        return []
    q = "trashed=false"
    if query:
        q += f" and name contains '{query.replace(chr(39), '')}'"

    all_raw: list[dict] = []

    # ── Step 1: 分頁抓 allDrives（含我的雲端硬碟 + 所有共用硬碟）──
    page_token = None
    page_size = min(max_results, 1000)
    fetched = 0
    while True:
        try:
            params = {
                "pageSize": page_size,
                "q": q,
                "orderBy": "modifiedTime desc",
                "fields": "nextPageToken,files(id,name,mimeType,size,modifiedTime,webViewLink,driveId)",
                "supportsAllDrives": True,
                "includeItemsFromAllDrives": True,
                "corpora": "allDrives",
            }
            if page_token:
                params["pageToken"] = page_token
            r = httpx.get(f"{DRIVE_API}/files",
                headers={"Authorization": f"Bearer {token}"},
                params=params, timeout=30)
            data = r.json()
            raw = data.get("files", [])
            all_raw.extend(raw)
            fetched += len(raw)
            page_token = data.get("nextPageToken")
            print(f"[drive_service] 已取 {fetched} 筆，nextPageToken={'有' if page_token else '無'}")
            if not page_token or fetched >= max_results:
                break
        except Exception as e:
            print(f"[drive_service] allDrives 搜尋失敗：{e}")
            break

    # ── Step 2: 取得共用硬碟名稱對照表，補充 drive_name 欄位 ──
    shared_drives = _fetch_shared_drives(token)
    drive_id_to_name = {d["id"]: d["name"] for d in shared_drives}

    for f in all_raw:
        drive_id = f.get("driveId", "")
        if drive_id and drive_id in drive_id_to_name:
            f["drive_name"] = drive_id_to_name[drive_id]
        else:
            f["drive_name"] = "我的雲端硬碟"

    # 去重（同一 id 只保留一筆）
    seen = set()
    deduped = []
    for f in all_raw:
        if f["id"] not in seen:
            seen.add(f["id"])
            deduped.append(f)

    print(f"[drive_service] 索引檔案數：{len(deduped)}，共用硬碟：{len(shared_drives)} 個")
    _save_to_index(db_func, deduped, user_conn=user_conn)
    return [{
        "id": f["id"],
        "name": f["name"],
        "type": _type_label(f.get("mimeType","")),
        "size": _fmt_size(f.get("size",0)),
        "modified": f.get("modifiedTime","")[:10],
        "url": f.get("webViewLink",""),
        "drive_name": f.get("drive_name","我的雲端硬碟"),
    } for f in deduped]

_SYNONYM_MAP = {
    "評估": ["評價", "鑑價", "估值", "估算", "評定"],
    "評價": ["評估", "鑑價", "估值"],
    "鑑價": ["評估", "評價", "估值"],
    "合約": ["合同", "協議書", "合作協議", "委任", "委託", "服務合約", "採購合約", "契約"],
    "合同": ["合約", "協議書"],
    "協議書": ["合約", "合作協議", "同意書"],
    "保密": ["NDA", "機密", "不外洩", "切結書"],
    "NDA": ["保密", "機密", "切結書"],
    "報價": ["估價", "報價單", "請款", "費用"],
    "估價": ["報價", "報價單"],
    "損益": ["盈虧", "損益表", "P&L", "財損"],
    "盈虧": ["損益", "損益表"],
    "財務": ["財報", "損益", "資產負債", "帳務", "會計"],
    "薪資": ["薪酬", "薪水", "薪金", "人事費用"],
    "薪酬": ["薪資", "薪水"],
    "預算": ["費用", "支出", "budget"],
    "業績": ["銷售", "營收", "業績統計", "Sales"],
    "提案": ["企劃", "方案", "簡報", "proposal"],
    "企劃": ["提案", "方案", "計畫"],
    "簡報": ["PPT", "pptx", "presentation", "提案"],
    "計畫": ["規劃", "方案", "plan"],
    "請款": ["請款單", "帳務", "報價"],
    "到貨": ["進貨", "收貨", "貨到"],
    "股權": ["股份", "出資", "持股"],
    "資產": ["固定資產", "財產", "資產負債"],
    "發票": ["收據", "統一發票"],
    "出貨": ["出倉", "發貨", "GD出貨"],
    "申報": ["申請", "報告", "申報書"],
}

def _fuzzy_tokens(query: str) -> list[str]:
    """拆出查詢關鍵詞，加入同義詞，去掉停用詞。"""
    import re as _re
    stop = {"阿福", "幫我", "找", "一下", "那個", "那份", "這份", "請", "可以", "文件",
            "檔案", "資料", "在哪", "在哪裡", "的", "了", "嗎", "呢", "啊", "幫",
            "看", "讀", "給我", "念", "摘要", "重點", "告訴我", "說", "找一下",
            "雲端", "Drive", "drive", "Google", "google", "硬碟", "本機", "Mac"}
    raw = _re.findall(r"[A-Za-z0-9_.-]{2,}|[一-鿿]{2,}", query)
    tokens = set()
    for t in raw:
        if t in stop:
            continue
        tokens.add(t)
        for syn in _SYNONYM_MAP.get(t, []):
            tokens.add(syn)
    return list(tokens)

def _score_drive_file(name: str, drive_name: str, tokens: list[str]) -> int:
    """給 Drive 檔案打分：token 命中越多越高，完全包含加分，drive_name 命中加分。"""
    combined = (name + " " + (drive_name or "")).lower()
    score = 0
    for tok in tokens:
        tl = tok.lower()
        if tl in combined:
            score += len(tok) * 10
            if tl in name.lower():
                score += len(tok) * 5  # 在檔名裡加分
    return score

def _read_from_index(db_func, query: str = "", limit: int = 20, user_conn=None) -> list[dict]:
    """Read from local SQLite index，支援模糊多關鍵詞 + 同義詞。"""
    c, should_close = _get_conn(user_conn)
    _ensure_drive_index_table(c)
    if query:
        tokens = _fuzzy_tokens(query)
        if tokens:
            # 對每個 token 做 OR LIKE，然後在 Python 端打分排序
            conditions = " OR ".join(["name LIKE ?" for _ in tokens])
            params = [f"%{t}%" for t in tokens]
            rows = c.execute(
                f"SELECT id,name,mime_type,size,modified,url,drive_name FROM drive_index "
                f"WHERE {conditions} ORDER BY modified DESC LIMIT 200",
                params
            ).fetchall()
            # 也加上 drive_name 搜尋
            drive_conditions = " OR ".join(["drive_name LIKE ?" for _ in tokens])
            drive_rows = c.execute(
                f"SELECT id,name,mime_type,size,modified,url,drive_name FROM drive_index "
                f"WHERE {drive_conditions} ORDER BY modified DESC LIMIT 100",
                params
            ).fetchall()
            # 合併去重
            seen = set()
            merged = []
            for r in rows + drive_rows:
                if r[0] not in seen:
                    seen.add(r[0])
                    merged.append(r)
            # 打分排序
            scored = sorted(
                merged,
                key=lambda r: _score_drive_file(r[1], r[6] or "", tokens),
                reverse=True
            )
            rows = scored[:limit]
        else:
            kw = f"%{query}%"
            rows = c.execute(
                "SELECT id,name,mime_type,size,modified,url,drive_name FROM drive_index "
                "WHERE name LIKE ? ORDER BY modified DESC LIMIT ?",
                (kw, limit)
            ).fetchall()
    else:
        rows = c.execute(
            "SELECT id,name,mime_type,size,modified,url,drive_name FROM drive_index ORDER BY modified DESC LIMIT ?",
            (limit,)
        ).fetchall()
    if should_close:
        c.close()
    return [{
        "id": r[0], "name": r[1], "type": _type_label(r[2]),
        "size": r[3], "modified": r[4], "url": r[5],
        "drive_name": r[6] or "我的雲端硬碟",
    } for r in rows]

def search_files(db_func, query: str = "", limit: int = 20, force_refresh: bool = False, user_conn=None) -> tuple[list[dict], bool]:
    """
    Return (files, from_cache).
    Reads from local index if fresh; fetches from API and rebuilds index otherwise.
    """
    if not force_refresh and _index_is_fresh(db_func, user_conn=user_conn):
        results = _read_from_index(db_func, query, limit, user_conn=user_conn)
        return results, True
    # Fetch fresh from API (slow, but builds index)
    results = _fetch_from_api(db_func, query, max_results=100, user_conn=user_conn)
    if query:
        q_lower = query.lower()
        results = [f for f in results if q_lower in f["name"].lower() or q_lower in f["type"].lower()]
    return results[:limit], False

def list_recent(db_func, limit: int = 10, user_conn=None) -> tuple[list[dict], bool]:
    return search_files(db_func, query="", limit=limit, user_conn=user_conn)

def index_count(db_func, user_conn=None) -> int:
    try:
        c, should_close = _get_conn(user_conn)
        _ensure_drive_index_table(c)
        n = c.execute("SELECT COUNT(*) FROM drive_index").fetchone()[0]
        if should_close:
            c.close()
        return n
    except Exception:
        return 0

def is_connected(db_func) -> bool:
    return _token(db_func) is not None


def _extract_pdf_bytes(content: bytes) -> str:
    import io, os as _os, subprocess as _subprocess, tempfile as _tempfile
    from pathlib import Path as _Path
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        if text and len(text.strip()) > 40:
            return text[:80000]
    except Exception:
        pass

    pdf_path = None
    try:
        with _tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pf:
            pf.write(content)
            pdf_path = pf.name

        try:
            with _tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tf:
                txt_path = tf.name
            try:
                _subprocess.run(["pdftotext", "-layout", pdf_path, txt_path], check=False, capture_output=True, timeout=45)
                with open(txt_path, encoding="utf-8", errors="ignore") as fh:
                    text = fh.read()
                if text and len(text.strip()) > 40:
                    return text[:80000]
            finally:
                try: _os.unlink(txt_path)
                except Exception: pass
        except Exception:
            pass

        with _tempfile.TemporaryDirectory(prefix="alfred_drive_pdf_ocr_") as td:
            prefix = _os.path.join(td, "page")
            _subprocess.run(["pdftoppm", "-r", "150", "-png", "-f", "1", "-l", "3", pdf_path, prefix], check=False, capture_output=True, timeout=90)
            chunks = []
            for img in sorted(_Path(td).glob("page-*.png"))[:3]:
                out = _subprocess.run(["tesseract", str(img), "stdout", "-l", "chi_tra+chi_sim+eng", "--psm", "6"], check=False, capture_output=True, text=True, timeout=90)
                if out.stdout and out.stdout.strip():
                    chunks.append(out.stdout.strip())
            return "\n\n".join(chunks)[:80000]
    except Exception as exc:
        return f"[OCR 失敗: {exc}]"
    finally:
        if pdf_path:
            try: _os.unlink(pdf_path)
            except Exception: pass
    return ""


def _ensure_text_cache_table(conn):
    conn.execute("""CREATE TABLE IF NOT EXISTS drive_text_cache (
        file_id TEXT PRIMARY KEY,
        mime_type TEXT,
        text TEXT,
        extracted_at TEXT
    )""")
    conn.commit()


def _get_cached_text(file_id: str, user_conn=None) -> str:
    try:
        c, should_close = _get_conn(user_conn)
        _ensure_text_cache_table(c)
        row = c.execute("SELECT text FROM drive_text_cache WHERE file_id=? LIMIT 1", (file_id,)).fetchone()
        if should_close:
            c.close()
        return row[0] if row and row[0] else ""
    except Exception:
        return ""


def _save_cached_text(file_id: str, mime_type: str, text: str, user_conn=None):
    if not text or len(text.strip()) < 40:
        return
    try:
        c, should_close = _get_conn(user_conn)
        _ensure_text_cache_table(c)
        c.execute(
            "INSERT OR REPLACE INTO drive_text_cache (file_id,mime_type,text,extracted_at) VALUES (?,?,?,?)",
            (file_id, mime_type or "", text[:80000], datetime.now().isoformat())
        )
        c.commit()
        if should_close:
            c.close()
    except Exception:
        pass

def download_and_extract(file_id: str, token: str, mime_type: str = "") -> str:
    """下載 Drive 檔案並抽取文字內容。支援 Google Docs export、一般檔案、PDF/OCR，並把結果快取。"""
    mime_type = mime_type or ""
    cached = _get_cached_text(file_id)
    if cached and len(cached.strip()) > 40:
        return cached

    headers = {"Authorization": f"Bearer {token}"}
    export_map = {
        "application/vnd.google-apps.document": "text/plain",
        "application/vnd.google-apps.spreadsheet": "text/csv",
        "application/vnd.google-apps.presentation": "text/plain",
    }

    try:
        text = ""
        if mime_type in export_map:
            r = httpx.get(
                f"{DRIVE_API}/files/{file_id}/export",
                headers=headers,
                params={"mimeType": export_map[mime_type]},
                timeout=30,
            )
            if r.status_code >= 400:
                return f"[匯出失敗: HTTP {r.status_code} {r.text[:200]}]"
            text = r.text[:80000]
        else:
            r = httpx.get(
                f"{DRIVE_API}/files/{file_id}",
                headers=headers,
                params={"alt": "media"},
                timeout=30,
            )
            if r.status_code >= 400:
                return f"[下載失敗: HTTP {r.status_code} {r.text[:200]}]"
            content = r.content
            # 圖片/ZIP/壓縮檔/二進位 → 直接跳過，不嘗試解析
            if any(t in mime_type for t in ["image/", "zip", "rar", "7z", "octet-stream", "tar"]) \
               or any(mime_type.endswith(e) for e in [".jpg", ".jpeg", ".png", ".gif", ".zip", ".rar"]):
                text = f"[NO_TEXT: 此檔案為{mime_type}，無法直接朗讀文字內容。]"
            elif "pdf" in mime_type:
                text = _extract_pdf_bytes(content)
            elif "wordprocessing" in mime_type or mime_type.endswith("docx") or mime_type.endswith(".doc"):
                import io
                try:
                    import docx
                    doc = docx.Document(io.BytesIO(content))
                    parts = [p.text for p in doc.paragraphs if p.text]
                    for table in doc.tables:
                        for row in table.rows:
                            vals = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                            if vals:
                                parts.append(" | ".join(vals))
                    text = "\n".join(parts)[:80000]
                except Exception:
                    text = ""
            elif "spreadsheet" in mime_type or mime_type.endswith("xlsx") or mime_type.endswith("xls") or "excel" in mime_type or "ms-excel" in mime_type:
                import io
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
                    parts = []
                    for sheet in wb.worksheets[:5]:
                        parts.append(f"[工作表: {sheet.title}]")
                        for row in sheet.iter_rows(max_row=200, values_only=True):
                            vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                            if vals:
                                parts.append("\t".join(vals))
                    text = "\n".join(parts)[:80000]
                except Exception:
                    try:
                        import xlrd, io as _io
                        wb = xlrd.open_workbook(file_contents=content)
                        parts = []
                        for sheet in wb.sheets()[:5]:
                            parts.append(f"[工作表: {sheet.name}]")
                            for rx in range(min(200, sheet.nrows)):
                                vals = [str(sheet.cell_value(rx, cx)).strip() for cx in range(sheet.ncols) if str(sheet.cell_value(rx, cx)).strip()]
                                if vals:
                                    parts.append("\t".join(vals))
                        text = "\n".join(parts)[:80000]
                    except Exception:
                        text = ""
            elif mime_type.endswith("msword") or mime_type == "application/msword":
                # Old .doc format — try docx fallback then plain
                import io
                try:
                    import docx
                    doc = docx.Document(io.BytesIO(content))
                    text = "\n".join(p.text for p in doc.paragraphs if p.text)[:80000]
                except Exception:
                    text = content.decode("utf-8", errors="ignore")[:80000]
            else:
                try:
                    text = content.decode("utf-8", errors="ignore")[:80000]
                except Exception:
                    text = ""
        if not text or len(text.strip()) < 40:
            text = "[NO_TEXT: 已嘗試抽取文字，但此檔案目前沒有足夠可朗讀內容。可能是掃描品質不足、圖片型 PDF，或檔案沒有文字層。]"
        _save_cached_text(file_id, mime_type, text)
        return text
    except Exception as e:
        return f"[下載失敗: {e}]"

